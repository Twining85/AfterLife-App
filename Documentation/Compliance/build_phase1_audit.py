from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "Documentation" / "Compliance" / "AfterLife_Phase_1_Ist_Audit.docx"

BLUE = "2E6F73"
DARK = "173F43"
LIGHT_BLUE = "E7F1F2"
LIGHT_GRAY = "F2F4F5"
RED = "9B1C1C"
AMBER = "8A5A00"
GREEN = "246B43"
MUTED = "5C6668"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Seite ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def add_field_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        set_cell_shading(cells[0], LIGHT_GRAY)
        cells[0].paragraphs[0].runs[0].bold = True
    set_table_widths(table, [2100, 7260])
    return table


def add_status_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, BLUE)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    set_repeat_table_header(table.rows[0])
    for row_values in rows:
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for idx, value in enumerate(row_values):
            cells[idx].text = str(value)
            if idx == 0 and str(value).startswith("P0"):
                set_cell_shading(cells[idx], "F8D7DA")
            elif idx == 0 and str(value).startswith("P1"):
                set_cell_shading(cells[idx], "FFF0CE")
            elif idx == 0 and str(value).startswith("P2"):
                set_cell_shading(cells[idx], LIGHT_BLUE)
    set_table_widths(table, widths)
    return table


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.add_run(text)
    return p


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.45)
section.footer_distance = Inches(0.45)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string("202526")
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

for style_name, size, color, before, after in (
    ("Title", 25, DARK, 0, 6),
    ("Subtitle", 13, MUTED, 0, 18),
    ("Heading 1", 16, BLUE, 18, 9),
    ("Heading 2", 13, BLUE, 13, 7),
    ("Heading 3", 11.5, DARK, 10, 5),
):
    style = styles[style_name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = style_name != "Subtitle"
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

for list_name in ("List Bullet", "List Bullet 2", "List Number"):
    styles[list_name].font.name = "Calibri"
    styles[list_name].font.size = Pt(10.5)
    styles[list_name].paragraph_format.space_after = Pt(4)
    styles[list_name].paragraph_format.line_spacing = 1.15

header = section.header
p = header.paragraphs[0]
p.text = "TSCHLÜSSLI  |  Datenschutz- und Sicherheitsprüfung"
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in p.runs:
    run.font.name = "Calibri"
    run.font.size = Pt(8.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(MUTED)
add_page_number(section.footer.paragraphs[0])

doc.add_paragraph("PHASE 1 · IST-AUDIT", style="Title")
doc.add_paragraph("Datenschutz, Informationssicherheit und Apple-Review-Bereitschaft", style="Subtitle")
add_field_table(doc, [
    ("Produkt", "Tschlüssli / internes Xcode-Projekt AfterLife"),
    ("Auditstand", "17. Juli 2026"),
    ("Prüfart", "Statische Prüfung von Quellcode und Projektkonfiguration"),
    ("Projektpfad", "/Users/reneengeler/AfterLife"),
    ("Status", "Erstbefund – noch keine Korrekturen umgesetzt"),
])

doc.add_heading("Zweck und Aussagekraft", level=1)
doc.add_paragraph(
    "Dieser Bericht dokumentiert den technischen Ist-Zustand der iOS-App anhand des lokal vorliegenden "
    "Quellcodes. Er trennt bestätigte Code-Fakten, daraus abgeleitete Risiken und extern zu klärende Punkte. "
    "Er ist Arbeits- und Abnahmedokument für die folgenden Umsetzungsphasen."
)
callout = doc.add_table(rows=1, cols=1)
callout.style = "Table Grid"
callout.cell(0, 0).text = (
    "Wichtige Grenze: Nicht geprüft wurden produktiver Netzwerkverkehr, Server- und Vercel-Konfiguration, "
    "Webseiteninhalte, Verträge mit Dienstleistern, App Store Connect, reale Geräte-Backups oder ein "
    "eingereichtes App-Archiv. Rechtliche Bewertungen sind durch eine qualifizierte Fachperson zu bestätigen."
)
set_cell_shading(callout.cell(0, 0), "FFF4D6")
set_table_widths(callout, [9360])

doc.add_heading("Management-Zusammenfassung", level=1)
doc.add_paragraph(
    "Der aktuelle Stand ist noch nicht freigabereif für die Verarbeitung der vorgesehenen hochsensiblen "
    "Vorsorge-, Gesundheits-, Finanz- und Zugangsdaten. Besonders kritisch ist, dass Passwörter trotz "
    "vorhandenem Keychain zusätzlich im Klartext in AppStorage und SwiftData geführt und in Exportausgaben "
    "übernommen werden. Außerdem werden IBANs an einen externen Dienst in der URL übertragen und "
    "Einladungstokens protokolliert beziehungsweise persistent in UserDefaults gehalten."
)
add_status_table(doc, ["Priorität", "Anzahl", "Bedeutung"], [
    ("P0", "4", "Vor produktiver Nutzung zwingend beheben"),
    ("P1", "8", "Vor App-Store-Freigabe beziehungsweise Pilotbetrieb beheben oder verbindlich klären"),
    ("P2", "4", "Härtung und dokumentierte Qualitätsverbesserung"),
    ("Positiv", "6", "Bereits vorhandene geeignete Grundlagen"),
], [1500, 1200, 6660])

doc.add_heading("Prioritätsdefinition", level=2)
bullet(doc, "P0 – unmittelbares Risiko für vertrauliche Daten oder grundlegende Sicherheitsarchitektur.")
bullet(doc, "P1 – wesentliche Apple-, Datenschutz- oder Sicherheitslücke vor Veröffentlichung.")
bullet(doc, "P2 – sinnvolle Härtung, Nachweisführung oder Wartbarkeitsverbesserung.")
bullet(doc, "OFFEN – im lokalen Code nicht abschließend feststellbar; externer Nachweis erforderlich.")

doc.add_heading("Kritische Befunde", level=1)
critical_rows = [
    ("P0-01", "Passwörter mehrfach im Klartext gespeichert", "`gespeichertesPasswort` liegt in AppStorage; `registrierungsPasswort` und Konto-Passwörter liegen in SwiftData. Damit wird der Keychain-Schutz umgangen.", "Registrierung.swift:44,790; ProfilModell.swift:31–34; Profil.swift:1045–1048; AboModell.swift:54–66"),
    ("P0-02", "Passwörter werden in Dossier/PDF-Ausgaben übernommen", "Profil- und Konto-Passwörter werden beim Zeichnen des Dossiers ausgegeben. Ein geteiltes PDF kann dadurch Zugangsdaten offenlegen.", "Profil.swift:2762, 2918, 2970–2974"),
    ("P0-03", "Vollständige IBAN an Drittanbieter übertragen", "Die IBAN wird als URL-Pfad an openiban.com gesendet. URLs können in Infrastruktur-, Proxy- oder Serverlogs landen; Rechtsgrundlage und Auftragsverarbeitung sind nicht belegt.", "Finanzen.swift:2083–2096"),
    ("P0-04", "Einladungstokens in Logs und UserDefaults", "Der vollständige Deep Link und Token werden ausgegeben und persistent gespeichert. Tokens gewähren potentiell Dossierzugriff und sind wie Geheimnisse zu behandeln.", "AfterLifeApp.swift:82–86, 267–307"),
]
add_status_table(doc, ["ID", "Befund", "Risiko", "Code-Nachweis"], critical_rows, [1000, 2200, 3460, 2700])

doc.add_heading("Wesentliche Befunde vor Freigabe", level=1)
p1_rows = [
    ("P1-01", "Privacy Manifest fehlt", "Im Projekt wurde keine `PrivacyInfo.xcprivacy` gefunden. Required-Reason-APIs und App-Datenangaben sind noch systematisch zu bestimmen.", "Projektdateien / Dateiinventar"),
    ("P1-02", "Temporäre sensible Dateien", "PDFs, Scans, Videos und Dokumente werden im temporären Verzeichnis geschrieben; konsistente Löschung und explizite Data-Protection-Optionen sind nicht nachgewiesen.", "Dokumente.swift:975–1007; Wuensche.swift:1862–1942; Profil.swift:1836–1839"),
    ("P1-03", "Rechtliche Links nicht eindeutig", "In der Registrierung ist nur ‚Nutzungsbedingungen‘ klickbar; ‚Datenschutz‘ ist Text. Im Profil öffnen rechtliche Punkte nur die Domain-Startseite.", "Registrierung.swift:504–520; Profil.swift:579–639"),
    ("P1-04", "Berechtigungstexte inkonsistent", "App-Name und Zwecke sind nicht einheitlich. Foto-Schreibzugriff nennt Videos, während auch Fotoalben gespeichert werden; Lesezugriff spricht nur vom Album.", "project.pbxproj:480–485, 522–527"),
    ("P1-05", "Lokales Passwortmodell statt sicherer Authentifizierung", "Mindestens sechs Zeichen, Vergleich mit lokalem Klartext und keine nachgewiesene serverseitige Verifikation/Hashing. Für ein echtes Konto ungeeignet.", "Profil.swift:1117–1173; Relogin.swift:338–354"),
    ("P1-06", "Externe Datenflüsse nicht belegt", "Adressdaten gehen an Vercel/Post und OpenPLZ; IBAN an OpenIBAN. Datenschutzbedingungen, Speicherfristen, Serverstandorte und Verträge liegen nicht im Projekt vor.", "Profil.swift:1188–1316; Finanzen.swift:2083–2096"),
    ("P1-07", "Löschung deckt temporäre Exporte nicht nachweislich ab", "SwiftData, Keychain, ausgewählte UserDefaults und Erinnerungen werden gelöscht; zuvor erzeugte temporäre Dateien werden dabei nicht gezielt bereinigt.", "Profil.swift:1319–1416"),
    ("P1-08", "Schutz der SwiftData-Datenbank nicht explizit nachgewiesen", "Die persistente ModelConfiguration verwendet die Standardablage. Eine bewusst geprüfte Schutzklasse, Backup-Strategie und Verschlüsselungsdokumentation fehlen.", "AfterLifeApp.swift:14–49"),
]
add_status_table(doc, ["ID", "Befund", "Bewertung", "Code-Nachweis"], p1_rows, [1000, 2200, 3460, 2700])

doc.add_heading("Härtungs- und Nachweispunkte", level=1)
p2_rows = [
    ("P2-01", "Produktionslogs", "Dateipfade, Termine und interne IDs werden per `print` ausgegeben. Release-Logging sollte datensparsam und zentral gesteuert werden.", "Dokumente.swift:1096–1120; NotificationService.swift:21–71; Profil.swift:994"),
    ("P2-02", "Universal Links", "Associated Domain ist konfiguriert; korrekte AASA-Datei, Team-ID/Bundle-ID und Tokenbehandlung auf tschluessli.ch müssen live validiert werden.", "AfterLife.entitlements; AfterLifeApp.swift:310–356"),
    ("P2-03", "App-Umschalter und Bildschirmaufnahme", "Kein expliziter Schutz sensibler Ansichten gegen App-Snapshots oder eine dokumentierte Entscheidung dazu gefunden.", "Kein Treffer im statischen Audit"),
    ("P2-04", "Automatisierte Sicherheitstests", "Vorhandene Tests enthalten keine erkennbare Abnahme für vollständige Datenlöschung, Exportbereinigung oder Secret-Leakage.", "AfterLifeTests / AfterLifeUITests"),
]
add_status_table(doc, ["ID", "Punkt", "Bewertung", "Nachweis"], p2_rows, [1000, 2200, 3460, 2700])

doc.add_heading("Bereits geeignete Grundlagen", level=1)
positive_rows = [
    ("POS-01", "Keychain-Zugriffsklasse", "`kSecAttrAccessibleWhenUnlockedThisDeviceOnly` verhindert Migration auf andere Geräte und schützt bei gesperrtem Gerät.", "KeychainHelper.swift:29–40, 82–98"),
    ("POS-02", "Verschlüsselte Netzwerktransporte", "Im statischen Audit wurden produktive HTTP-Aufrufe ausschließlich mit HTTPS gefunden; keine ATS-Ausnahme erkennbar.", "Profil.swift; Finanzen.swift"),
    ("POS-03", "Bestätigte Profillöschung", "Explizite destruktive Bestätigung sowie Löschung vieler SwiftData-Modelle, Keychain-Einträge, Zustände und Erinnerung vorhanden.", "Profil.swift:610–624, 1319–1416"),
    ("POS-04", "Mitteilungsberechtigung im Nutzungskontext", "Berechtigung wird im Erinnerungsablauf angefragt und Ablehnung kann zu den Einstellungen führen.", "NotificationService.swift; Home.swift:671–678"),
    ("POS-05", "System-Fotoauswahl", "PhotosPicker reduziert unnötigen pauschalen Zugriff beim Import einzelner Medien.", "Dokumente.swift:442; Wuensche.swift:784,849"),
    ("POS-06", "Keine externen Swift Packages", "Im Xcode-Projekt sind aktuell keine Package-Produkte eingebunden; dadurch ist die Drittanbieter-SDK-Fläche klein.", "project.pbxproj:174,197,220"),
]
add_status_table(doc, ["ID", "Grundlage", "Bewertung", "Code-Nachweis"], positive_rows, [1000, 2200, 3460, 2700])

doc.add_heading("Vorläufiges Dateninventar", level=1)
doc.add_paragraph(
    "Die folgende Übersicht ist aus den SwiftData-Modellen und Nutzungsstellen abgeleitet. Sie ist in Phase 2 "
    "um Empfänger, Rechtsgrundlage, Aufbewahrungsfrist und endgültige Schutzmaßnahme zu ergänzen."
)
inventory_rows = [
    ("Profil und Identität", "Name, Geburtstag, Adresse, Telefon, E-Mail, AHV-Nummer, Profilbild", "SwiftData; teilweise AppStorage", "Sehr hoch"),
    ("Authentifizierung", "E-Mail, Passwort, Loginstatus, biometrischer Status", "Keychain, AppStorage, SwiftData", "Sehr hoch"),
    ("Gesundheit", "Hausarzt, Blutgruppe, Allergien, Medikamente, Organspende, Hinweise", "SwiftData", "Sehr hoch"),
    ("Finanzen", "Bankkonten, IBAN, Schulden, Versicherungen, Immobilien, Wertsachen, Steuerunterlagen", "SwiftData; OpenIBAN", "Sehr hoch"),
    ("Digitale Konten", "Anbieter, Benutzername, Passwort, Geräte-PIN, Notizen", "SwiftData", "Sehr hoch"),
    ("Vorsorge und Wünsche", "Bestattung, letzte Botschaft, Patientenverfügung, Vorsorgeauftrag, Krankheit", "SwiftData/externer Speicher", "Sehr hoch"),
    ("Dokumente und Medien", "Scans, Fotos, Videos, Testamente, PDFs", "SwiftData/externer Speicher/temp. Dateien", "Sehr hoch"),
    ("Drittpersonen", "Hinterbliebene, Ärzte, Vertrauenspersonen, Kontaktdaten", "SwiftData", "Hoch"),
    ("Freigaben", "Dossier-IDs, E-Mail, Rollen, Status, Einladungs- und Zugriffstokens", "SwiftData, UserDefaults, URL", "Sehr hoch"),
    ("Nutzungszustand", "Navigation, Bereichsstatus, Erinnerungsdatum, Berechtigungsstatus", "AppStorage/UserDefaults", "Mittel"),
]
add_status_table(doc, ["Kategorie", "Beispiele", "Speicher/Empfänger", "Schutzbedarf"], inventory_rows, [1900, 3500, 2460, 1500])

doc.add_heading("Externe Dienste und Datenflüsse", level=1)
external_rows = [
    ("Vercel-Proxy / Schweizerische Post", "Straßensuche und Gebäudeprüfung", "Suchtext; bei Verifikation Straße, Hausnummer, PLZ, Ort", "Vertrag, Logs, Standort, Löschfrist und Datenschutzhinweis offen"),
    ("OpenPLZ API", "Ortsbestimmung zur PLZ", "Postleitzahl", "Datenschutzhinweis und Serverprotokollierung offen"),
    ("OpenIBAN", "IBAN-Prüfung", "Vollständige IBAN im URL-Pfad", "P0; vor Nutzung ersetzen oder verbindlich absichern"),
    ("Frankfurter API", "Wechselkurse", "Währungscodes; keine erkennbaren Personendaten", "Geringes Datenschutzrisiko; Verfügbarkeit dokumentieren"),
    ("tschluessli.ch", "Rechtliches und Einladungslinks", "Einladungstoken in Query-Parameter", "Webserver-Logs, Ablauf, Widerruf und Einmalnutzung klären"),
]
add_status_table(doc, ["Dienst", "Zweck", "Übermittelte Daten", "Offene Prüfung"], external_rows, [2100, 2100, 2800, 2360])

doc.add_heading("Funktionsbezogene Prüfergebnisse", level=1)
doc.add_heading("Berechtigungen", level=2)
bullet(doc, "Kamera: Usage Description vorhanden; tatsächlichen Scan-Start und Ablehnungsfall auf realem Gerät prüfen.")
bullet(doc, "Face ID: Usage Description vorhanden; Text ist nicht konsistent mit dem Store-Namen Tschlüssli.")
bullet(doc, "Fotomediathek: selektiver PhotosPicker beim Import ist positiv; direkte Add-/ReadWrite-Abfragen beim Export müssen mit korrekten Texten abgeglichen werden.")
bullet(doc, "Mitteilungen: kontextbezogener Ablauf vorhanden; Sperrbildschirmtext auf sensible Inhalte prüfen.")

doc.add_heading("Löschung", level=2)
bullet(doc, "In-App-Löschoption ist leicht auffindbar und verlangt eine klare Bestätigung.")
bullet(doc, "Viele registrierte SwiftData-Modelle werden gelöscht; Keychain-Konten und die jährliche Erinnerung werden entfernt.")
bullet(doc, "Temporäre Exporte, Vorschauen und möglicherweise bereits geteilte Dateien können technisch nicht durch die bestehende Routine nachweislich erfasst werden.")
bullet(doc, "Serverseitige Datenlöschung ist mangels prüfbarer Backend-Implementierung offen.")

doc.add_heading("Authentifizierung und Biometrie", level=2)
bullet(doc, "Keychain-Implementierung nutzt eine angemessene Zugriffsklasse.")
bullet(doc, "Der Sicherheitsgewinn wird durch parallele Klartextkopien in AppStorage und SwiftData aufgehoben.")
bullet(doc, "Face ID wird über LocalAuthentication verwendet; der bekannte Lifecycle-Sonderfall wird im Routing berücksichtigt.")
bullet(doc, "Ob überhaupt ein echtes serverseitiges Konto besteht, muss als Produktentscheidung geklärt werden.")

doc.add_heading("Dokumente und Export", level=2)
bullet(doc, "Dokumente und Videos werden als Daten in SwiftData beziehungsweise externalStorage abgelegt.")
bullet(doc, "Für Vorschau und Teilen werden Klartextdateien im temporären Verzeichnis erzeugt.")
bullet(doc, "Dateinamen werden nur teilweise bereinigt; Schutzklasse, Lebensdauer und Löschung sind nicht zentralisiert.")
bullet(doc, "Das Dossier kann Passwörter ausgeben; dies ist vor jeder externen Testverteilung zu entfernen.")

doc.add_heading("Offene Nachweise außerhalb des Repositories", level=1)
open_items = [
    "Aktuelle Datenschutzerklärung mit direkter HTTPS-URL und Versionsstand",
    "Nutzungsbedingungen und Impressum mit direkter URL",
    "App-Store-Privacy-Antworten und gegebenenfalls Tracking-Angaben",
    "Vercel-/Backend-Quellcode, Logging, Secrets, Zugriffsschutz und Löschkonzept",
    "Verträge beziehungsweise Datenschutzbedingungen aller externen API-Anbieter",
    "Serverstandorte und grenzüberschreitende Datenübermittlungen",
    "AASA-Datei und Live-Funktion der Universal Links",
    "Backup-Verhalten und Data-Protection-Klasse auf einem realen Gerät",
    "Netzwerkmitschnitt eines vollständigen Testablaufs",
    "TestFlight/App-Archiv einschließlich Privacy Report und Validierungsprotokoll",
]
for item in open_items:
    bullet(doc, item)

doc.add_heading("Empfohlene Reihenfolge für Phase 2 und 3", level=1)
steps = [
    ("1", "Secret-Sofortmaßnahme", "Klartextpasswörter aus AppStorage, Profilmodellen, digitalen Konten und Exporten entfernen; Datenmigration definieren."),
    ("2", "Token- und Logging-Härtung", "Tokens nicht protokollieren; sichere Ablage, Ablauf, Einmalnutzung und Widerruf festlegen."),
    ("3", "IBAN-Datenfluss stoppen", "OpenIBAN-Aufruf bis zur Architektur- und Vertragsentscheidung deaktivieren oder lokal/über kontrolliertes Backend ersetzen."),
    ("4", "Dateischutz zentralisieren", "Temporäre Dateien mit Schutzklasse, eindeutiger Lebensdauer und zentraler Bereinigung verwalten."),
    ("5", "Datenschutzoberfläche", "Direkte Datenschutz-, Impressums- und Nutzungsbedingungen-Links sowie verständliche Permission-Texte einbauen."),
    ("6", "Manifest und Store-Angaben", "Privacy Manifest, Required-Reason-API-Prüfung und App-Store-Datenschutzlabel erstellen."),
    ("7", "Externe Nachweise", "Dienstleister, Verträge, Serverlogs, Löschfristen und Auslandsübermittlungen klären."),
    ("8", "Geräte- und Archivtest", "Data Protection, Backup, Netzwerkverkehr, Löschung und Xcode-Archiv praktisch validieren."),
]
add_status_table(doc, ["Schritt", "Paket", "Ergebnis"], steps, [900, 2400, 6060])

doc.add_heading("Abnahmekriterien für Phase 1", level=1)
criteria = [
    ("Erfüllt", "Prüfumfang und Grenzen dokumentiert"),
    ("Erfüllt", "Datenkategorien und externe Dienste vorläufig inventarisiert"),
    ("Erfüllt", "Befunde mit Priorität und Code-Nachweis dokumentiert"),
    ("Erfüllt", "Positive Kontrollen und bestehende Grundlagen dokumentiert"),
    ("Erfüllt", "Offene externe Nachweise separat ausgewiesen"),
    ("Offen", "Bericht durch Product Owner bestätigt"),
    ("Offen", "Start der Sofortmaßnahmen P0 freigegeben"),
]
add_status_table(doc, ["Status", "Kriterium"], criteria, [1500, 7860])

doc.add_heading("Quellenbasis", level=1)
doc.add_paragraph("Interne Nachweise")
bullet(doc, "Xcode-Projekt und Swift-Dateien unter /Users/reneengeler/AfterLife, statisch geprüft am 17. Juli 2026.")
doc.add_paragraph("Offizielle Apple-Grundlagen")
for source in [
    "Apple App Review Guidelines: https://developer.apple.com/app-store/review/guidelines/",
    "Apple – Manage app privacy: https://developer.apple.com/help/app-store-connect/manage-app-information/manage-app-privacy",
    "Apple – Offering account deletion: https://developer.apple.com/support/offering-account-deletion-in-your-app",
    "Apple – Privacy manifests: https://developer.apple.com/documentation/bundleresources/adding-a-privacy-manifest-to-your-app-or-third-party-sdk",
    "Apple – Using the Keychain: https://developer.apple.com/documentation/security/using-the-keychain-to-manage-user-secrets",
]:
    bullet(doc, source)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
